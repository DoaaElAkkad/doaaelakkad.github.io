import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_enhanced_cv():
    # Create a new Document
    doc = docx.Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add header with personal information
    header = doc.add_heading('DOAA ELAKKAD', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_format = header.runs[0].font
    header_format.size = Pt(18)
    header_format.bold = True
    header_format.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    
    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run('Head of Market Research | Cairo, Egypt 11837\n').bold = True
    contact_info.add_run('+201005875595 | doaa.elakkad88@gmail.com | www.linkedin.com/in/doaa-elakkad-55a86211b')
    
    # Add horizontal line
    add_horizontal_line(doc)
    
    # Add Professional Summary
    add_section_heading(doc, 'PROFESSIONAL SUMMARY')
    summary = doc.add_paragraph()
    summary.add_run('Results-driven Market Research Leader with 10+ years of progressive experience in strategic market analysis and commercial partnerships. Demonstrated expertise in designing and implementing comprehensive research methodologies to identify investment opportunities and market trends. Skilled in translating complex data into actionable insights for executive decision-making. Proven track record of leading high-performing research teams and developing data-driven strategies that enhance organizational performance and market positioning.')
    
    # Add Core Competencies
    add_section_heading(doc, 'CORE COMPETENCIES')
    skills = [
        'Strategic Market Analysis & Forecasting',
        'Qualitative & Quantitative Research Design',
        'Consumer Behavior & Trend Analysis',
        'Data Analytics & Statistical Modeling (SPSS)',
        'Project Management & Team Leadership',
        'Commercial Partnership Development',
        'Executive Reporting & Presentations',
        'Microsoft Office Suite (Advanced)',
        'Cross-functional Collaboration'
    ]
    
    # Create a 3-column table for skills
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.autofit = True
    
    # Fill the table with skills
    cell_index = 0
    for i in range(3):
        for j in range(3):
            if cell_index < len(skills):
                cell = table.cell(i, j)
                cell.text = "• " + skills[cell_index]
                cell_index += 1
    
    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            set_cell_border(cell, top=None, bottom=None, start=None, end=None)
    
    # Add Professional Experience
    add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # Head of Market Research at EgyptAir
    add_job_heading(doc, 'Head of Market Research', 'EgyptAir', 'Cairo, Egypt', 'July 2024 - Present')
    add_job_bullets(doc, [
        'Develop and implement comprehensive research strategies aligned with organizational objectives, resulting in data-driven decision-making across departments',
        'Lead a team of research professionals, establishing clear performance metrics and fostering a culture of analytical excellence',
        'Design innovative research methodologies to capture market trends, consumer preferences, and competitive insights',
        'Deliver executive-level presentations that translate complex data into strategic recommendations, directly influencing corporate planning'
    ])
    
    # Market Research Expert at Star Alliance
    add_job_heading(doc, 'Market Research Expert', 'Star Alliance, MRE Member', 'Cairo, Egypt', 'July 2024 - Present')
    add_job_bullets(doc, [
        'Spearhead the development of sophisticated Net Promoter Score (NPS) questionnaires, enhancing customer satisfaction measurement across alliance partners',
        'Ensure statistical validity and reliability of research methodologies through rigorous quota validation processes',
        'Synthesize complex multi-market data into comprehensive reports that drive strategic decision-making at the alliance level'
    ])
    
    # Head of Market Research at Broketopia
    add_job_heading(doc, 'Head of Market Research', 'Broketopia for Real Estate', 'Cairo, Egypt', 'September 2021 - January 2023')
    add_job_bullets(doc, [
        'Established and implemented research best practices that increased data quality by 35% and reduced research cycle time by 20%',
        'Led a cross-functional team of 8 researchers, developing talent and improving team productivity by 25%',
        'Conducted comprehensive competitive analysis that identified 3 key market opportunities, directly contributing to 15% revenue growth',
        'Directed UX/UI research for mobile applications, resulting in a 40% improvement in user engagement metrics',
        'Presented actionable insights to C-suite executives, directly influencing product development and go-to-market strategies'
    ])
    
    # Senior Market Research at EgyptAir Airlines
    add_job_heading(doc, 'Senior Market Research', 'EgyptAir Airlines', 'Cairo, Egypt', 'January 2019 - August 2021')
    add_job_bullets(doc, [
        'Designed and executed mixed-method research studies that identified key consumer segments and purchasing behaviors, informing targeted marketing campaigns',
        'Developed and administered customer satisfaction surveys using SurveyMonkey, achieving a 28% response rate (industry average: 20%)',
        'Managed a mystery shopping program across 12 locations, improving service quality metrics by 22% within 6 months',
        'Conducted concept testing for 5 new service offerings, identifying the 2 most viable options that were successfully launched',
        'Analyzed complex datasets using SPSS and Excel, generating insights that informed a 15% increase in customer retention'
    ])
    
    # Mystery Shopper at G.W.R Consulting
    add_job_heading(doc, 'Mystery Shopper', 'G.W.R Consulting', 'Cairo, Egypt', 'August 2021 - July 2022')
    add_job_bullets(doc, [
        'Conducted over 50 mystery shopping evaluations across multiple industries, providing detailed qualitative and quantitative feedback',
        'Developed comprehensive reports that identified specific service improvement opportunities, contributing to client service enhancement initiatives',
        'Maintained 100% compliance with evaluation protocols and reporting deadlines'
    ])
    
    # Market Research & Strategic Partnerships Specialist at EgyptAir Holding Co.
    add_job_heading(doc, 'Market Research & Strategic Partnerships Specialist', 'EgyptAir Holding Co.', 'Cairo, Egypt', 'January 2015 - December 2019')
    add_job_bullets(doc, [
        'Analyzed global aviation market trends using Lufthansa DDS and other intelligence systems, identifying 5 high-potential growth markets',
        'Evaluated and secured 7 strategic partnership opportunities that expanded network reach by 15% and increased passenger revenue by 12%',
        'Negotiated and implemented codeshare agreements with 4 major international carriers, enhancing global market presence',
        'Developed comprehensive partnership management frameworks, ensuring seamless integration of operational, technical, and commercial requirements',
        'Created and delivered strategic presentations to executive leadership, securing approval for 85% of proposed partnership initiatives'
    ])
    
    # Accountant at EgyptAir Maintenance and Engineering
    add_job_heading(doc, 'Accountant', 'EgyptAir Maintenance and Engineering', 'Cairo, Egypt', 'January 2011 - December 2015')
    add_job_bullets(doc, [
        'Conducted cost-benefit analyses for maintenance services, identifying pricing optimizations that increased profit margins by 8%',
        'Collaborated with marketing department to develop competitive pricing strategies for maintenance service offerings',
        'Performed competitive market analyses that informed strategic decisions on service partnerships, resulting in 3 new high-value contracts'
    ])
    
    # Customer Service Agent at Vodafone Egypt
    add_job_heading(doc, 'Customer Service Agent for Market Research', 'Vodafone Egypt', 'Cairo, Egypt', 'January 2010 - December 2011')
    add_job_bullets(doc, [
        'Conducted over 1,000 customer satisfaction surveys, achieving a 90% completion rate',
        'Collected and organized qualitative feedback on new product offerings, contributing to product refinement and marketing strategy'
    ])
    
    # Add Education
    add_section_heading(doc, 'EDUCATION')
    
    # MBA
    education1 = doc.add_paragraph()
    education1.add_run('Master of Business Administration: Marketing\n').bold = True
    education1.add_run('Arab Academy for Science, Technology & Maritime | 2018 - 2020')
    
    # Bachelor's
    education2 = doc.add_paragraph()
    education2.add_run('Bachelor of Commerce\n').bold = True
    education2.add_run('Cairo University | 2005 - 2009')
    
    # Add Professional Development
    add_section_heading(doc, 'PROFESSIONAL DEVELOPMENT')
    
    certifications = [
        'Learning Excel: Data Analytics (March 2025)',
        'An Intuitive Introduction to Probabilities, University of Zurich (February 2025)',
        'Marketing Foundations: Analytics Languages (January 2025)',
        'Professional Digital Marketing, Udacity FWD Scholarship (June 2021)',
        'Mystery Shopping Training, G.W.R Consulting (January 2021)',
        'Foundation Certificate in Marketing, American University in Cairo (March 2018)',
        'International Negotiation Skills, Arab Air Carriers Organization (February 2018)',
        'Certificate of Achievement in Marketing and Sales, American University in Cairo (December 2017)',
        'Advanced Sales and Distribution Strategies, Arab Air Carriers Organization (July 2017)',
        'Fundamental Marketing Workshop, Core Management (January 2015)',
        'Basic Business Skills Acquisition, Future Generation Foundation (July 2010)'
    ]
    
    # Create a 2-column table for certifications
    cert_table = doc.add_table(rows=6, cols=2)
    cert_table.style = 'Table Grid'
    cert_table.autofit = True
    
    # Fill the table with certifications
    cell_index = 0
    for i in range(6):
        for j in range(2):
            if cell_index < len(certifications):
                cell = cert_table.cell(i, j)
                cell.text = "• " + certifications[cell_index]
                cell_index += 1
    
    # Remove table borders
    for row in cert_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            set_cell_border(cell, top=None, bottom=None, start=None, end=None)
    
    # Add Languages
    add_section_heading(doc, 'LANGUAGES')
    languages = doc.add_paragraph()
    languages.add_run('Arabic (Native) | English (Professional Working Proficiency)')
    
    # Save the document
    doc.save('/home/ubuntu/cv_project/Doaa_Elakkad_Enhanced_CV.docx')
    print("Enhanced CV created successfully!")

def add_section_heading(doc, heading_text):
    heading = doc.add_paragraph()
    heading.add_run(heading_text).bold = True
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    
    # Add bottom border to the paragraph
    p = heading._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_job_heading(doc, title, company, location, dates):
    job = doc.add_paragraph()
    job.add_run(title + '\n').bold = True
    company_line = job.add_run(company + ' | ' + location + ' | ' + dates)
    company_line.italic = True

def add_job_bullets(doc, bullet_points):
    for point in bullet_points:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(point)
        bullet.paragraph_format.left_indent = Inches(0.25)
    
    # Add a small space after the bullet points
    doc.add_paragraph()

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()
    p = p._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    """
    Set cell border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    tcBorders = OxmlElement('w:tcBorders')
    
    if top is not None:
        t_border = OxmlElement('w:top')
        t_border.set(qn('w:val'), 'nil' if top is None else 'single')
        t_border.set(qn('w:sz'), '4')
        t_border.set(qn('w:space'), '0')
        t_border.set(qn('w:color'), 'auto')
        tcBorders.append(t_border)
    
    if bottom is not None:
        b_border = OxmlElement('w:bottom')
        b_border.set(qn('w:val'), 'nil' if bottom is None else 'single')
        b_border.set(qn('w:sz'), '4')
        b_border.set(qn('w:space'), '0')
        b_border.set(qn('w:color'), 'auto')
        tcBorders.append(b_border)
    
    if start is not None:
        s_border = OxmlElement('w:start')
        s_border.set(qn('w:val'), 'nil' if start is None else 'single')
        s_border.set(qn('w:sz'), '4')
        s_border.set(qn('w:space'), '0')
        s_border.set(qn('w:color'), 'auto')
        tcBorders.append(s_border)
    
    if end is not None:
        e_border = OxmlElement('w:end')
        e_border.set(qn('w:val'), 'nil' if end is None else 'single')
        e_border.set(qn('w:sz'), '4')
        e_border.set(qn('w:space'), '0')
        e_border.set(qn('w:color'), 'auto')
        tcBorders.append(e_border)
    
    if tcBorders.getchildren():
        tcPr.append(tcBorders)

if __name__ == "__main__":
    create_enhanced_cv()
